from src.core.format_reporting import format_reporting
from src.core.abstract_report import abstract_report
from src.core.custom_raise import CustomRaise

from docx import Document
import os

class docx_report(abstract_report):
    def __init__(self) -> None:
        super().__init__()
        self.__format = format_reporting.DOCX
    
    def create(self, data: list):
        CustomRaise.type_exception("data", data, list)
        if len(data) == 0:
            CustomRaise.operation_exception("Набор данных пуст")

        self.document = Document()
        first_model = data[0]

        if isinstance(first_model, tuple):
            fields = [f"column_{i}" for i in range(len(first_model))]
        else:
            fields = list(filter(lambda x: not x.startswith("_") and not callable(getattr(first_model, x)), dir(first_model)))

        for row in data:
            for field in fields:
                value = getattr(row, field)
                if isinstance(value, list):
                    for item in value:
                        self.document.add_paragraph(f"  - {str(item)}")
                else:
                    self.document.add_paragraph(f"{field.capitalize()}: {str(value)}")
            self.document.add_paragraph("\n")

    def save(self, directory: str, filename: str):
        full_path = os.path.abspath(os.path.join(os.path.dirname(__file__), directory))
        
        if not os.path.exists(full_path):
            os.makedirs(full_path)
        
        try:
            file_path = os.path.join(full_path, filename + ".docx")
            self.document.save(file_path)
            return True
        except:
            return False